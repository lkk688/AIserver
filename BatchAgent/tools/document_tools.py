"""
放这些内容：
	•	DocumentToolManager

职责：
	•	解析 PDF 路径
	•	缓存 document agent
	•	load document
	•	get_overview
	•	read_section
	•	search_document

这是 UniversalToolHandler 和 document_agent_v2 之间的桥。

.cache/documents/
  sources/
    <doc_key>.pdf
  extracted/
    <doc_key>/
      document.md
      meta.json
      ocr_workspace/
      assets/   # 后续你若让 extractor 把 assets 定向到这里
  rag/
    <doc_key>/
      chunks.json
      rag_meta.json
"""

from __future__ import annotations

import hashlib
import json
import os
import re
import shutil
import subprocess
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Dict, Optional, Tuple

from BatchAgent.ocr_agent.document_agent import (
    create_document_agent_from_pdf,
    DocumentAgent,
)


# Optional imports
try:
    from bs4 import BeautifulSoup
except ImportError:
    BeautifulSoup = None

try:
    import docx
except ImportError:
    docx = None


SUPPORTED_EXTENSIONS = {
    ".pdf",
    ".docx",
    ".doc",
    ".txt",
    ".md",
    ".markdown",
    ".html",
    ".htm",
}


@dataclass
class DocumentSourceInfo:
    source_type: str            # local_file
    original_path: str
    local_path: Path
    format: str                # pdf/docx/doc/txt/md/html
    identity: Dict[str, Any]
    doc_key: str


class DocumentToolManager:
    """
    Local multi-format document manager.

    Features:
    - local file only
    - dedicated document cache dir (not session dir)
    - in-memory cache + disk cache
    - unified API for overview / section read / search

    Cache layout:
      <document_cache_dir>/
        extracted/
          <doc_key>/
            document.md
            meta.json
            assets/
        converted/
          <doc_key>.pdf         # optional for DOC -> PDF conversion
    """

    def __init__(self, config: Any):
        self.config = config
        self.loaded_agents: Dict[str, Any] = {}
        self.loaded_source_infos: Dict[str, DocumentSourceInfo] = {}
        self.active_doc_key: Optional[str] = None

        self.doc_cache_root = Path(
            getattr(config, "document_cache_dir", ".cache/documents")
        ).resolve()

        self.extracted_cache_dir = self.doc_cache_root / "extracted"
        self.converted_cache_dir = self.doc_cache_root / "converted"

        for d in [self.doc_cache_root, self.extracted_cache_dir, self.converted_cache_dir]:
            d.mkdir(parents=True, exist_ok=True)

    # ------------------------------------------------------------------
    # Source resolution
    # ------------------------------------------------------------------

    def _project_root(self) -> Path:
        return Path(__file__).resolve().parent.parent.parent

    def _resolve_local_path(self, filepath: str) -> Optional[Path]:
        raw = Path(filepath)
        search_order = [
            raw,
            Path.cwd() / raw,
            self._project_root() / raw,
        ]
        for candidate in search_order:
            if candidate.exists() and candidate.is_file():
                return candidate.resolve()
        return None

    def _detect_format(self, path: Path) -> str:
        ext = path.suffix.lower()
        if ext not in SUPPORTED_EXTENSIONS:
            raise ValueError(
                f"Unsupported document format: '{ext}'. "
                f"Supported: {sorted(SUPPORTED_EXTENSIONS)}"
            )
        return ext.lstrip(".")

    def _compute_hash_key(self, identity: Dict[str, Any]) -> str:
        payload = json.dumps(identity, sort_keys=True, ensure_ascii=False).encode("utf-8")
        return hashlib.sha1(payload).hexdigest()[:16]

    def _resolve_document_source(self, filepath: str) -> DocumentSourceInfo:
        local_path = self._resolve_local_path(filepath)
        if local_path is None:
            raise FileNotFoundError(
                f"File not found: '{filepath}'. "
                "Provide a valid local document path."
            )

        fmt = self._detect_format(local_path)
        stat = local_path.stat()
        identity = {
            "type": "local_file",
            "path": str(local_path),
            "size": stat.st_size,
            "mtime": stat.st_mtime,
            "format": fmt,
        }
        doc_key = self._compute_hash_key(identity)

        return DocumentSourceInfo(
            source_type="local_file",
            original_path=filepath,
            local_path=local_path,
            format=fmt,
            identity=identity,
            doc_key=doc_key,
        )

    # ------------------------------------------------------------------
    # Cache paths
    # ------------------------------------------------------------------

    def _get_cache_paths(self, doc_key: str) -> Dict[str, Path]:
        extract_dir = self.extracted_cache_dir / doc_key
        assets_dir = extract_dir / "assets"
        converted_pdf = self.converted_cache_dir / f"{doc_key}.pdf"

        return {
            "extract_dir": extract_dir,
            "assets_dir": assets_dir,
            "markdown_file": extract_dir / "document.md",
            "meta_file": extract_dir / "meta.json",
            "converted_pdf": converted_pdf,
        }

    def _disk_cache_exists(self, source_info: DocumentSourceInfo) -> bool:
        cache_paths = self._get_cache_paths(source_info.doc_key)
        return cache_paths["markdown_file"].exists() and cache_paths["meta_file"].exists()

    def _write_extract_meta(
        self,
        source_info: DocumentSourceInfo,
        markdown_file: Path,
    ):
        cache_paths = self._get_cache_paths(source_info.doc_key)
        cache_paths["extract_dir"].mkdir(parents=True, exist_ok=True)

        meta = {
            "doc_key": source_info.doc_key,
            "source_type": source_info.source_type,
            "original_path": source_info.original_path,
            "local_path": str(source_info.local_path),
            "format": source_info.format,
            "markdown_file": str(markdown_file),
            "identity": source_info.identity,
        }

        cache_paths["meta_file"].write_text(
            json.dumps(meta, indent=2, ensure_ascii=False),
            encoding="utf-8",
        )

    # ------------------------------------------------------------------
    # Embedding mode
    # ------------------------------------------------------------------

    def _choose_embedding_mode(self) -> str:
        return "api" if os.environ.get("EMBEDDING_BASE_URL") else "local"

    # ------------------------------------------------------------------
    # Markdown builders for non-PDF formats
    # ------------------------------------------------------------------

    def _clean_text(self, text: str) -> str:
        text = text.replace("\r\n", "\n").replace("\r", "\n")
        text = re.sub(r"[ \t]+", " ", text)
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text.strip()

    def _normalize_plain_text_to_markdown(self, text: str) -> str:
        text = self._clean_text(text)
        if not text:
            return ""

        lines = text.splitlines()
        out = []
        current_para = []

        heading_patterns = [
            re.compile(r"^(abstract|introduction|background|related work|method|methods|methodology|approach|experiments|results|discussion|conclusion|conclusions|references|appendix)\.?$", re.I),
            re.compile(r"^(\d+(\.\d+)*|[A-Z]|[IVXLC]+)[\.\)]?\s+[A-Z].*$"),
        ]

        def flush_para():
            nonlocal current_para
            if current_para:
                out.append(" ".join(s.strip() for s in current_para if s.strip()))
                current_para = []

        for line in lines:
            s = line.strip()
            if not s:
                flush_para()
                continue

            is_heading = any(p.match(s) for p in heading_patterns)
            if is_heading and len(s.split()) <= 14:
                flush_para()
                out.append(f"## {s.rstrip('.').strip()}")
                continue

            current_para.append(s)

        flush_para()
        return "\n\n".join(out).strip()

    def _extract_txt_or_md_to_markdown(self, source_info: DocumentSourceInfo) -> str:
        text = source_info.local_path.read_text(encoding="utf-8", errors="replace")
        if source_info.format in {"md", "markdown"}:
            return self._clean_text(text)
        return self._normalize_plain_text_to_markdown(text)

    def _extract_html_to_markdown(self, source_info: DocumentSourceInfo) -> str:
        html = source_info.local_path.read_text(encoding="utf-8", errors="replace")

        if BeautifulSoup is None:
            # fallback: crude strip
            text = re.sub(r"<script[\s\S]*?</script>", " ", html, flags=re.I)
            text = re.sub(r"<style[\s\S]*?</style>", " ", text, flags=re.I)
            text = re.sub(r"<[^>]+>", "\n", text)
            return self._normalize_plain_text_to_markdown(text)

        soup = BeautifulSoup(html, "html.parser")

        for tag in soup(["script", "style", "noscript"]):
            tag.decompose()

        out = []
        for el in soup.find_all(["h1", "h2", "h3", "h4", "h5", "h6", "p", "li"]):
            text = el.get_text(" ", strip=True)
            if not text:
                continue

            if el.name.startswith("h"):
                level = int(el.name[1])
                level = min(max(level, 1), 6)
                out.append(f"{'#' * level} {text}")
            elif el.name == "li":
                out.append(f"- {text}")
            else:
                out.append(text)

        return self._clean_text("\n\n".join(out))

    def _extract_docx_to_markdown(self, source_info: DocumentSourceInfo) -> str:
        if docx is None:
            raise RuntimeError(
                "python-docx is not installed. Install it to process .docx files."
            )

        document = docx.Document(str(source_info.local_path))
        out = []

        for para in document.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            style_name = (para.style.name or "").lower() if para.style else ""

            if "heading" in style_name:
                m = re.search(r"heading\s*(\d+)", style_name)
                if m:
                    level = int(m.group(1))
                    level = min(max(level, 1), 6)
                else:
                    level = 2
                out.append(f"{'#' * level} {text}")
            else:
                out.append(text)

        # tables
        for t_idx, table in enumerate(document.tables, start=1):
            rows = []
            for row in table.rows:
                rows.append([cell.text.strip().replace("\n", " ") for cell in row.cells])

            if rows and rows[0]:
                header = rows[0]
                md = []
                md.append("| " + " | ".join(header) + " |")
                md.append("| " + " | ".join(["---"] * len(header)) + " |")
                for row in rows[1:]:
                    md.append("| " + " | ".join(row) + " |")
                out.append(f"**Table {t_idx}**\n\n" + "\n".join(md))

        return self._clean_text("\n\n".join(out))

    def _convert_doc_to_pdf(self, source_info: DocumentSourceInfo) -> Path:
        """
        Convert legacy .doc to PDF using LibreOffice headless.
        """
        cache_paths = self._get_cache_paths(source_info.doc_key)
        out_pdf = cache_paths["converted_pdf"]
        if out_pdf.exists():
            return out_pdf

        out_pdf.parent.mkdir(parents=True, exist_ok=True)

        tmp_out_dir = out_pdf.parent
        cmd = [
            "libreoffice",
            "--headless",
            "--convert-to",
            "pdf",
            "--outdir",
            str(tmp_out_dir),
            str(source_info.local_path),
        ]

        try:
            subprocess.run(cmd, check=True, capture_output=True)
        except FileNotFoundError:
            raise RuntimeError(
                "libreoffice not found. Install LibreOffice to process .doc files."
            )
        except subprocess.CalledProcessError as e:
            raise RuntimeError(
                f"Failed to convert DOC to PDF: {e.stderr.decode('utf-8', errors='replace')}"
            )

        converted_name = source_info.local_path.with_suffix(".pdf").name
        converted_path = tmp_out_dir / converted_name
        if not converted_path.exists():
            raise RuntimeError("DOC conversion completed but output PDF not found.")

        # normalize to stable cache filename
        shutil.move(str(converted_path), str(out_pdf))
        return out_pdf

    def _extract_non_pdf_to_markdown(self, source_info: DocumentSourceInfo) -> str:
        fmt = source_info.format

        if fmt in {"txt", "md", "markdown"}:
            return self._extract_txt_or_md_to_markdown(source_info)

        if fmt in {"html", "htm"}:
            return self._extract_html_to_markdown(source_info)

        if fmt == "docx":
            return self._extract_docx_to_markdown(source_info)

        if fmt == "doc":
            # Convert DOC -> PDF, then reuse PDF pipeline by building fresh from converted file
            converted_pdf = self._convert_doc_to_pdf(source_info)
            converted_info = DocumentSourceInfo(
                source_type="local_file",
                original_path=source_info.original_path,
                local_path=converted_pdf,
                format="pdf",
                identity={**source_info.identity, "converted_from": "doc"},
                doc_key=source_info.doc_key,
            )
            agent = self._create_agent_fresh_pdf(converted_info, converted_pdf)
            return getattr(agent, "raw_text", "")

        raise ValueError(f"Unsupported non-PDF format: {fmt}")

    # ------------------------------------------------------------------
    # Agent creation
    # ------------------------------------------------------------------

    def _create_agent_fresh_pdf(self, source_info: DocumentSourceInfo, pdf_path: Path):
        emb_type = self._choose_embedding_mode()
        cache_paths = self._get_cache_paths(source_info.doc_key)
        extract_dir = cache_paths["extract_dir"]
        extract_dir.mkdir(parents=True, exist_ok=True)

        ocr_workspace = str(extract_dir / "ocr_workspace")

        agent = create_document_agent_from_pdf(
            pdf_filepath=str(pdf_path),
            ocr_server=getattr(self.config, "ocr_server", None),
            ocr_model=getattr(self.config, "ocr_model", "allenai/olmOCR-2-7B-1025-FP8"),
            ocr_workspace=ocr_workspace,
            use_pdf_page_ocr=getattr(self.config, "use_pdf_page_ocr", False),
            embedding_type=emb_type,
            local_model_name=getattr(self.config, "local_embedding_model", "all-MiniLM-L6-v2"),
            api_url=os.environ.get("EMBEDDING_BASE_URL", ""),
            api_key=os.environ.get("EMBEDDING_API_KEY", "EMPTY"),
            api_model=os.environ.get("EMBEDDING_MODEL", ""),
        )

        markdown_text = getattr(agent, "raw_text", None)
        if markdown_text:
            cache_paths["extract_dir"].mkdir(parents=True, exist_ok=True)
            cache_paths["markdown_file"].write_text(markdown_text, encoding="utf-8")

        self._write_extract_meta(source_info=source_info, markdown_file=cache_paths["markdown_file"])
        return agent

    def _create_agent_fresh_non_pdf(self, source_info: DocumentSourceInfo):
        emb_type = self._choose_embedding_mode()
        cache_paths = self._get_cache_paths(source_info.doc_key)
        cache_paths["extract_dir"].mkdir(parents=True, exist_ok=True)

        markdown_text = self._extract_non_pdf_to_markdown(source_info)
        cache_paths["markdown_file"].write_text(markdown_text, encoding="utf-8")
        self._write_extract_meta(source_info=source_info, markdown_file=cache_paths["markdown_file"])

        agent = DocumentAgent(
            markdown_text=markdown_text,
            embedding_type=emb_type,
            local_model_name=getattr(self.config, "local_embedding_model", "all-MiniLM-L6-v2"),
            api_url=os.environ.get("EMBEDDING_BASE_URL", ""),
            api_key=os.environ.get("EMBEDDING_API_KEY", "EMPTY"),
            api_model=os.environ.get("EMBEDDING_MODEL", ""),
            auto_build_rag=True,
        )
        return agent

    def _create_agent_from_disk_cache(self, source_info: DocumentSourceInfo):
        cache_paths = self._get_cache_paths(source_info.doc_key)
        markdown_file = cache_paths["markdown_file"]

        if not markdown_file.exists():
            raise FileNotFoundError(f"Cached markdown not found: {markdown_file}")

        markdown_text = markdown_file.read_text(encoding="utf-8")
        emb_type = self._choose_embedding_mode()

        agent = DocumentAgent(
            markdown_text=markdown_text,
            embedding_type=emb_type,
            local_model_name=getattr(self.config, "local_embedding_model", "all-MiniLM-L6-v2"),
            api_url=os.environ.get("EMBEDDING_BASE_URL", ""),
            api_key=os.environ.get("EMBEDDING_API_KEY", "EMPTY"),
            api_model=os.environ.get("EMBEDDING_MODEL", ""),
            auto_build_rag=True,
        )
        return agent

    # ------------------------------------------------------------------
    # Public API
    # ------------------------------------------------------------------

    def load_document(self, filepath: str) -> str:
        try:
            source_info = self._resolve_document_source(filepath)
            doc_key = source_info.doc_key
            self.loaded_source_infos[doc_key] = source_info

            # 1. memory cache
            if doc_key in self.loaded_agents:
                self.active_doc_key = doc_key
                return f"Loaded document from memory cache: {source_info.local_path.name}"

            # 2. disk cache
            if self._disk_cache_exists(source_info):
                agent = self._create_agent_from_disk_cache(source_info)
                self.loaded_agents[doc_key] = agent
                self.active_doc_key = doc_key
                return f"Loaded document from disk cache: {source_info.local_path.name}"

            # 3. fresh processing
            if source_info.format == "pdf":
                agent = self._create_agent_fresh_pdf(source_info, source_info.local_path)
            else:
                agent = self._create_agent_fresh_non_pdf(source_info)

            self.loaded_agents[doc_key] = agent
            self.active_doc_key = doc_key
            return f"Processed and cached document: {source_info.local_path.name}"

        except Exception as e:
            return f"Failed to load document '{filepath}': {e}"

    def get_active_agent(self):
        if self.active_doc_key and self.active_doc_key in self.loaded_agents:
            return self.loaded_agents[self.active_doc_key]
        return None

    def get_overview(self, filepath: Optional[str] = None) -> str:
        if filepath:
            msg = self.load_document(filepath)
            if msg.startswith("Failed to load document"):
                return msg

        agent = self.get_active_agent()
        if not agent:
            return (
                "No document is currently loaded. "
                "Call get_document_overview with a filepath first."
            )

        try:
            return agent.get_overview()
        except Exception as e:
            return f"Error getting document overview: {e}"

    def read_section(self, section_id: Optional[str] = None) -> str:
        agent = self.get_active_agent()
        if not agent:
            return (
                "No document is currently loaded. "
                "Call get_document_overview with a filepath first."
            )

        if not section_id:
            return "Please provide a section_id from get_document_overview."

        try:
            return agent.read_details(section_id=section_id)
        except Exception as e:
            return f"Error reading document section '{section_id}': {e}"

    def search_document(self, query: str, top_k: int = 5) -> str:
        agent = self.get_active_agent()
        if not agent:
            return (
                "No document is currently loaded. "
                "Call get_document_overview with a filepath first."
            )

        query = (query or "").strip()
        if not query:
            return "Please provide a non-empty search query."

        try:
            return agent.search(query=query, top_k=top_k)
        except Exception as e:
            return f"Error searching document: {e}"

    def get_loaded_documents(self) -> str:
        if not self.loaded_source_infos:
            return "No documents loaded."

        lines = ["Loaded documents:"]
        for doc_key, source_info in self.loaded_source_infos.items():
            marker = " [active]" if doc_key == self.active_doc_key else ""
            lines.append(
                f"- {source_info.original_path} -> doc_key={doc_key} format={source_info.format}{marker}"
            )
        return "\n".join(lines)

    def set_active_document(self, filepath: str) -> str:
        try:
            source_info = self._resolve_document_source(filepath)
            doc_key = source_info.doc_key
        except Exception as e:
            return f"Failed to resolve document '{filepath}': {e}"

        if doc_key not in self.loaded_agents:
            return (
                f"Document not loaded yet: '{filepath}'. "
                "Load it first with get_document_overview(filepath=...)."
            )

        self.active_doc_key = doc_key
        return f"Active document set to: {filepath}"

    def clear_memory_cache(self) -> str:
        self.loaded_agents.clear()
        self.loaded_source_infos.clear()
        self.active_doc_key = None
        return "Cleared in-memory document cache."

    def purge_document_cache(self, filepath: str) -> str:
        try:
            source_info = self._resolve_document_source(filepath)
            cache_paths = self._get_cache_paths(source_info.doc_key)

            if cache_paths["extract_dir"].exists():
                shutil.rmtree(cache_paths["extract_dir"])

            if cache_paths["converted_pdf"].exists():
                cache_paths["converted_pdf"].unlink()

            self.loaded_agents.pop(source_info.doc_key, None)
            self.loaded_source_infos.pop(source_info.doc_key, None)
            if self.active_doc_key == source_info.doc_key:
                self.active_doc_key = None

            return f"Purged cache for document: {filepath}"
        except Exception as e:
            return f"Failed to purge cache for '{filepath}': {e}"


# =============================================================================
# Test helpers
# =============================================================================

class _TestConfig:
    def __init__(self):
        self.document_cache_dir = ".cache/documents"
        self.ocr_server = None
        self.ocr_model = "allenai/olmOCR-2-7B-1025-FP8"
        self.use_pdf_page_ocr = False
        self.local_embedding_model = "all-MiniLM-L6-v2"


def test_resolve_and_key(path: str):
    cfg = _TestConfig()
    mgr = DocumentToolManager(cfg)
    info = mgr._resolve_document_source(path)

    print("=== test_resolve_and_key ===")
    print("source_type:", info.source_type)
    print("format:", info.format)
    print("original_path:", info.original_path)
    print("local_path:", info.local_path)
    print("doc_key:", info.doc_key)
    print("identity:", json.dumps(info.identity, indent=2, ensure_ascii=False))


def test_cache_paths(path: str):
    cfg = _TestConfig()
    mgr = DocumentToolManager(cfg)
    info = mgr._resolve_document_source(path)
    paths = mgr._get_cache_paths(info.doc_key)

    print("=== test_cache_paths ===")
    for k, v in paths.items():
        print(f"{k}: {v}")


def test_load_document(path: str):
    cfg = _TestConfig()
    mgr = DocumentToolManager(cfg)

    print("=== test_load_document (1st time) ===")
    print(mgr.load_document(path))
    print(mgr.get_loaded_documents())

    print("\n=== test_load_document (2nd time) ===")
    print(mgr.load_document(path))
    print(mgr.get_loaded_documents())


def test_overview(path: str):
    cfg = _TestConfig()
    mgr = DocumentToolManager(cfg)

    print("=== test_overview ===")
    print(mgr.get_overview(filepath=path))


def test_search(path: str, query: str):
    cfg = _TestConfig()
    mgr = DocumentToolManager(cfg)

    print("=== test_search ===")
    print(mgr.load_document(path))
    print(mgr.search_document(query=query, top_k=5))


if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser(description="DocumentToolManager local multi-format tests")
    parser.add_argument("--test", choices=["resolve", "paths", "load", "overview", "search"], default="resolve")
    parser.add_argument("--path", type=str, required=True, help="Local document path")
    parser.add_argument("--query", type=str, default="quantization stochastic rounding")

    args = parser.parse_args()

    if args.test == "resolve":
        test_resolve_and_key(args.path)
    elif args.test == "paths":
        test_cache_paths(args.path)
    elif args.test == "load":
        test_load_document(args.path)
    elif args.test == "overview":
        test_overview(args.path)
    elif args.test == "search":
        test_search(args.path, args.query)
"""
python -m BatchAgent.tools.document_tools \
  --test resolve \
  --path data/pdftest/NVIDIA-Nemotron-3-Super-Technical-Report.pdf

(py312) lkk@rtx5090:/Developer/AIserver$ python -m BatchAgent.tools.document_tools \
  --test paths \
  --path data/pdftest/NVIDIA-Nemotron-3-Super-Technical-Report.pdf


#really load data
python -m BatchAgent.tools.document_tools \
  --test load \
  --path data/pdftest/NVIDIA-Nemotron-3-Super-Technical-Report.pdf

python -m BatchAgent.tools.document_tools \
  --test overview \
  --path data/pdftest/NVIDIA-Nemotron-3-Super-Technical-Report.pdf

python -m BatchAgent.tools.document_tools \
  --test search \
  --path data/pdftest/NVIDIA-Nemotron-3-Super-Technical-Report.pdf

python -m BatchAgent.tools.document_tools \
  --test overview \
  --path data/markdown/data/pdftest/test3.md 
  
python -m BatchAgent.tools.document_tools \
  --test search \
  --path data/markdown/data/pdftest/test3.md  \
  --query "Q13"


pip install python-docx
python -m BatchAgent.tools.document_tools \
  --test overview \
  --path data/report.docx

pip install beautifulsoup4
python -m BatchAgent.tools.document_tools \
  --test overview \
  --path data/page.html
  
pip install libreoffice #process doc
"""
